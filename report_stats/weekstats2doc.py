#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time : 2020/7/13 9:18
# @Author : 马飞
# @File : report_stats.py.py
# @Software: PyCharm
import sys
import datetime
import pymysql
import json,os
from docx import Document
from docx.shared import Cm, Inches, Pt
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor
'''
  全局配置
'''
config = {
   'db_string'        : '10.2.39.18:3306:puppet:puppet:Puppet@123',
   'shape'            : 'line',
   'shape_disk_uage'  : 'bar',
   'start_date'       : '2020-07-11',
   'end_date'         : '2020-07-16',
   'cpu_title'        : 'cpu使用率[%]',
   'mem_title'        : '内存使用率7[%]',
   'disk_usage_title' : '磁盘使用率[%]',
   'disk_read_title'  : '磁盘读[KB]',
   'disk_write_title' : '磁盘写[KB]',
   'net_in_title'     : '网络流入[KB]',
   'net_out_title'    : '网络流出[KB]',
   'db_total_title'   : '总连接数',
   'db_active_title'  : '活跃连接数',
   'db_bak_sz_title'  : '备份大小[MB]',
   'db_bak_tm_title'  : '备份时间[S]',
   'cpu_color'        : '#8CD5C2',
   'mem_color'        : '#95CACA',
   'disk_usage_color' : '#5CADAD,#9CD5C4,#BCD5C5',
   'disk_read_color'  : '#8CD5C2',
   'disk_write_color' : '#8CD5C2',
   'net_in_color'     : '#8CD5C2',
   'net_out_color'    : '#8CD5C5',
   'db_total_color'   : '#8CD5C2',
   'db_active_color'  : '#8CD5C5',
   'db_bak_sz_color'  : '#8CD5C5',
   'db_bak_tm_color'  : '#8CD5C5',
}

'''
  功能：参与统计的服务器和数据库列表
'''
tj_server =[
     # 商管-备份服务器
     {
        'server_id': 81,
        'db_id': '',
     }
]



'''
   功能：获取数据库连接
'''
def get_cfg(cfg):
    db_ip      = cfg['db_string'].split(':')[0]
    db_port    = cfg['db_string'].split(':')[1]
    db_service = cfg['db_string'].split(':')[2]
    db_user    = cfg['db_string'].split(':')[3]
    db_pass    = cfg['db_string'].split(':')[4]
    cfg['db_mysql_dict'] = get_ds_mysql_dict(db_ip, db_port, db_service, db_user, db_pass)
    return cfg


def get_bzrq():
    return datetime.datetime.now().strftime("%Y年%m月%d日 %H时%M分%S秒")


'''
 功能：获取服务器信息
'''
def get_server(config):
    db=config['db_mysql_dict']
    cr=db.cursor()
    cr.execute('select * from t_server where id={}'.format(config['server_id']))
    rs=cr.fetchone()
    return rs

'''
 功能：获取数据库信息
'''
def get_db_server(config):
    db=config['db_mysql_dict']
    cr=db.cursor()
    cr.execute('select * from t_db_source where id={}'.format(config['db_id']))
    rs=cr.fetchone()
    return rs


'''
 功能:HTML模板
'''
html_templete='''
<!DOCTYPE html>
<html>
    <head>
         <meta charset="utf-8">
         <title>商管项目运行周报</title>
         <link href="https://cdn.bootcdn.net/ajax/libs/twitter-bootstrap/3.4.1/css/bootstrap.min.css" rel="stylesheet">
         <style>
           hr {
              -moz-border-bottom-colors: none;
              -moz-border-image: none;
              -moz-border-left-colors: none;
              -moz-border-right-colors: none;
              -moz-border-top-colors: none;
              border-color: #EEEEEE -moz-use-text-color #FFFFFF;
              border-style: solid none;
              border-width: 1px 0;
              margin: 18px 0;
            }
            ul {
              padding-bottom:30px;
            }
           .xwtable {width: 100%;border-collapse: collapse;border: 1px solid #ccc;}
           .xwtable thead td {font-size: 18px;color: #333333;
                              text-align: center;background: url(table_top.jpg) repeat-x top center;
                              border: 1px solid #ccc; font-weight:bold;}
           .xwtable thead th {font-size: 18px;color: #333333;
                              text-align: center;background: url(table_top.jpg) repeat-x top center;
                              border: 1px solid #ccc; font-weight:bold;}
           .xwtable tbody tr {background: #fff;font-size: 18px;color: #666666;}
           .xwtable tbody tr.alt-row {background: #f2f7fc;}
           .xwtable td{line-height:20px;text-align: left;padding:4px 10px 3px 10px;height: 18px;border: 1px solid #ccc;}
           span { color:red;}
            
         </style>
    </head>	
    <body>
       <div class="jumbotron jumbotron-fluid">
            <div class="container">
                  <h1>$$HEADER$$-运行周报</h1>                
                  <p>编制日期：$$BZRQ$$</p>
                  <p>编制人员：好生活平台组</p>
                  <p>统计日期：$$TJRQQ$$至$$TJRQZ$$</p>
            </div>
       </div>
       
       
       <!--运行情况 -->
       <div class="panel panel-default">
         <div class="panel-body">
                <ul class="nav nav-pills m-b-30 pull-left">
                    <li class="active">
                        <a href="#sg-week" data-toggle="tab" aria-expanded="true">本周运行情况</a>
                    </li>
                    <li class="">
                        <a href="#sg-server" data-toggle="tab" aria-expanded="true">服务器运行情况</a>
                    </li>
                    <li  class="">
                        <a href="#sg-db" data-toggle="tab" aria-expanded="false">$$DB_TAB$$</a>
                    </li>
                </ul>
                
                <div class="tab-content br-n pn">
                    <div id="sg-week" class="tab-pane active">
                         <div class="row">
                                <div class="col-md-12">
                                    <h3>$$SERVER_TITLE$$</h3>
                                    <div id='server_info' style="height:200px;">
                                       $$SERVER_TABLE$$
                                    </div>
                                 </div>
                           </div>
                           <div class="row">
                                 <div class="col-md-12">
                                    <h3>$$DB_TITLE$$</h3>
                                    <div id='db_info' style="height:200px;">
                                       $$DB_TABLE$$
                                    </div>
                                 </div>
                           </div> 
                    </div>
                    <div id="sg-server" class="tab-pane">
                          <div class="row">
                                <div class="col-md-12">
                                    <div id='cpu_usage' style="height:400px;"></div>
                                 </div>
                           </div>
                           <div class="row">
                                 <div class="col-md-12">
                                    <div id='mem_usage' style="height:400px;"></div>
                                 </div>
                           </div> 
                            <div class="row">
                                 <div class="col-md-12">
                                    <div id='disk_usage' style="height:400px;"></div>
                                 </div>
                           </div>                           
                           <div class="row">
                                 <div class="col-md-12">
                                    <div id='disk_read' style="height:400px;"></div>
                                 </div>
                           </div> 
                            <div class="row">
                                 <div class="col-md-12">
                                    <div id='disk_write' style="height:400px;"></div>
                                 </div>
                           </div>                           
                           <div class="row">
                                 <div class="col-md-12">
                                    <div id='net_in' style="height:400px;"></div>
                                 </div>
                           </div> 
                            <div class="row">
                                 <div class="col-md-12">
                                    <div id='net_out' style="height:400px;"></div>
                                 </div>
                           </div>
                    </div>
                    <div id="sg-db" class="tab-pane">
                           <div class="row">
                                <div class="col-md-12">
                                    <div id='db_total_num' style="height:400px;"></div>
                                 </div>
                           </div>
                           <div class="row">
                                 <div class="col-md-12">
                                    <div id='db_active_num' style="height:400px;"></div>
                                 </div>
                           </div> 
                           <div class="row">
                                <div class="col-md-12">
                                    <div id='db_bak_sz' style="height:400px;"></div>
                                 </div>
                           </div>
                           <div class="row">
                                 <div class="col-md-12">
                                    <div id='db_bak_tm' style="height:400px;"></div>
                                 </div>
                           </div>                           
                           
                    </div>
                </div>
          </div>
       </div>
      
       <script src="https://cdn.bootcdn.net/ajax/libs/jquery/3.5.1/jquery.js"></script>
       <script src="https://cdn.bootcdn.net/ajax/libs/twitter-bootstrap/3.4.1/js/bootstrap.min.js"></script>
       <script src="https://cdn.bootcdn.net/ajax/libs/echarts/4.8.0/echarts.min.js"></script>
       <script>
          function gen_charts(p_id,p_option) {
              var myChart = echarts.init($(p_id)[0]);
              var option  = p_option
              myChart.setOption(option);
          }
         
          $(document).ready(function() {
          
            $('a[data-toggle="tab"]').on('show.bs.tab', function (e) {
                  var activeTab = $(e.target).text();
                  var previousTab = $(e.relatedTarget).text();
                  console.log(activeTab,previousTab)
    
                  if (activeTab=='本周运行情况') {
                     $('#sg-week').show();
                     $('#sg-server').hide();
                     $('#sg-db').hide();
                  }
    
                  if (activeTab=='服务器运行情况') {
                     $('#sg-week').hide();
                     $('#sg-server').show();
                     $('#sg-db').hide();
                     gen_charts('#cpu_usage',  $cpu_option);
                     gen_charts('#mem_usage',  $mem_option);
                     gen_charts('#disk_usage', $disk_usage_option);
                     gen_charts('#disk_read',  $disk_read_option);
                     gen_charts('#disk_write', $disk_write_option);
                     gen_charts('#net_in',     $net_in_option);
                     gen_charts('#net_out',    $net_out_option);
                  }
                  
                  if (activeTab=='数据库运行情况') {
                     $('#sg-week').hide();
                     $('#sg-server').hide();
                     $('#sg-db').show();
                     gen_charts('#db_total_num',  $db_total_option);
                     gen_charts('#db_active_num', $db_active_option);
                     gen_charts('#db_bak_sz',     $db_bak_sz_option);
                     gen_charts('#db_bak_tm',     $db_bak_tm_option);
                  }
             });
            
          })
        </script>
    </body>
</html>
'''


'''
  功能：服务器概述
'''
st_server = '''
SELECT     
   'cpu使用率'                 AS flag,  
   ROUND(MIN(cpu_total_usage),2) AS val_min,
   ROUND(MAX(cpu_total_usage),2) AS val_max,
   ROUND(AVG(cpu_total_usage),2) AS val_avg,
   '阀值:连续3次>85%'             AS threshold,
   '未触发告警'                   AS message
FROM `t_monitor_task_server_log` 
WHERE server_id={}
AND  create_date between '{} 0:0:0' and '{} 23:59:59' 
UNION ALL
SELECT     
   '内存使用率' AS flag,  
   ROUND(MIN(mem_usage),2) AS val_min,
   ROUND(MAX(mem_usage),2) AS val_max,
   ROUND(AVG(mem_usage),2) AS val_avg,
   '阀值:连续3次>85%',
   '未触发告警' 
FROM `t_monitor_task_server_log` 
WHERE server_id={}
AND  create_date between '{} 0:0:0' and '{} 23:59:59' 
UNION ALL
SELECT     
   '磁盘使用率' AS flag,  
   MIN(disk_usage) AS val_min,
   MAX(disk_usage) AS val_max,
   '' AS val_avg,
   '阀值:连续3次>80%',
   '未触发告警'    
FROM `t_monitor_task_server_log` 
WHERE server_id={}
AND  create_date between '{} 0:0:0' and '{} 23:59:59' '''

st_db='''
SELECT     
   '总连接数' AS flag,  
   MIN(total_connect+0) AS val_min,
   MAX(total_connect+0) AS val_max,
   ROUND(AVG(total_connect+0),2) AS val_avg,
   '阀值 >300%'                AS threshold,
   '未触发告警'                AS message
FROM `t_monitor_task_db_log` 
WHERE db_id={}
AND  create_date between '{} 0:0:0' and '{} 23:59:59' 
UNION ALL
SELECT     
     '活跃连接数' AS flag,  
     MIN(active_connect+0) AS val_min,
     MAX(active_connect+0) AS val_max,
     ROUND(AVG(active_connect+0),2) AS val_avg,
    '阀值 >100%'                AS threshold,
    '未触发告警'                AS message
FROM `t_monitor_task_db_log` 
WHERE db_id={}
AND  create_date between '{} 0:0:0' and '{} 23:59:59' 
UNION ALL
SELECT 
    '备份大小[MB]' AS flag,
    MIN(CASE 
	    WHEN INSTR(total_size,'G')>0 THEN REPLACE(total_size,'G','')*1024
	    WHEN INSTR(total_size,'M')>0 THEN REPLACE(total_size,'M','')
	    ELSE REPLACE(total_size,'K','')/1024
	END) AS val_min,
    MAX(CASE 
	    WHEN INSTR(total_size,'G')>0 THEN REPLACE(total_size,'G','')*1024
	    WHEN INSTR(total_size,'M')>0 THEN REPLACE(total_size,'M','')
	    ELSE REPLACE(total_size,'K','')/1024
	END) AS val_max,
    ROUND(AVG(
	CASE 
	    WHEN INSTR(total_size,'G')>0 THEN REPLACE(total_size,'G','')*1024
	    WHEN INSTR(total_size,'M')>0 THEN REPLACE(total_size,'M','')
	    ELSE REPLACE(total_size,'K','')/1024
	END),2) AS val_avg,
	''    AS threshold,
	''    AS message
FROM `t_db_backup_total`
WHERE db_tag=(SELECT db_tag FROM t_db_config WHERE db_id={}) 
AND  create_date between '{} 0:0:0' and '{} 23:59:59' 
UNION ALL
SELECT 
    '备份时长[s]' AS flag,
    MIN(elaspsed_backup) AS val_min,
    MAX(elaspsed_backup) AS val_max,
    ROUND(AVG(elaspsed_backup),2) AS val_avg,
    ''    AS threshold,
    ''    AS message
FROM `t_db_backup_total`
WHERE db_tag=(SELECT db_tag FROM t_db_config WHERE db_id={}) 
AND  create_date between '{} 0:0:0' and '{} 23:59:59' 
'''


'''
  功能：服务器指标统计语句
'''
st_cpu        = '''SELECT  DATE_FORMAT(create_date,'%Y-%m-%d %H:%i') AS rq,
                           ROUND(AVG(cpu_total_usage),2) AS val 
                   FROM `t_monitor_task_server_log` 
                   WHERE server_id={}  
                     AND  create_date between '{} 0:0:0' and '{} 23:59:59' 
                   GROUP BY DATE_FORMAT(create_date,'%Y-%m-%d %H:%i') 
                   order by 1'''

st_mem        = '''SELECT  DATE_FORMAT(create_date,'%Y-%m-%d %H:%i') AS rq,
                           ROUND(AVG(mem_usage),2)  AS val 
                   FROM `t_monitor_task_server_log` 
                   WHERE server_id={}  
                      AND  create_date between '{} 0:0:0' and '{} 23:59:59' 
                   GROUP BY DATE_FORMAT(create_date,'%Y-%m-%d %H:%i') 
                   order by 1'''

st_disk_usage = '''SELECT DATE_FORMAT(create_date, '%Y-%m-%d') AS rq, 
                          disk_usage AS val 
                   FROM  `t_monitor_task_server_log` 
                   WHERE server_id = {} 
                      AND  create_date between '{} 0:0:0' and '{} 23:59:59' 
                   ORDER BY 1'''

st_disk_read  = '''SELECT  DATE_FORMAT(create_date,'%Y-%m-%d %H:%i') AS rq,
                           ROUND(AVG(disk_read/1024),2)  AS val 
                   FROM `t_monitor_task_server_log` 
                   WHERE server_id={}  
                     AND  create_date between '{} 0:0:0' and '{} 23:59:59' 
                   GROUP BY DATE_FORMAT(create_date,'%Y-%m-%d %H:%i') 
                   order by 1'''

st_disk_write = '''SELECT  DATE_FORMAT(create_date,'%Y-%m-%d %H:%i') AS rq,
                           ROUND(AVG(disk_write/1024),2)  AS val 
                   FROM `t_monitor_task_server_log` 
                   WHERE server_id={}  
                     AND  create_date between '{} 0:0:0' and '{} 23:59:59' 
                   GROUP BY DATE_FORMAT(create_date,'%Y-%m-%d %H:%i') 
                    order by 1'''

st_net_in   = '''SELECT  DATE_FORMAT(create_date,'%Y-%m-%d %H:%i') AS rq,
                           ROUND(AVG(net_in/1024),2)  AS val 
                 FROM `t_monitor_task_server_log` 
                 WHERE server_id={}  
                     AND  create_date between '{} 0:0:0' and '{} 23:59:59' 
                   GROUP BY DATE_FORMAT(create_date,'%Y-%m-%d %H:%i') 
                   order by 1'''

st_net_out = '''SELECT  DATE_FORMAT(create_date,'%Y-%m-%d %H:%i') AS rq,
                        ROUND(AVG(net_out/1024),2)  AS val 
                FROM `t_monitor_task_server_log` 
                 WHERE server_id={}  
                 AND  create_date between '{} 0:0:0' and '{} 23:59:59' 
                GROUP BY DATE_FORMAT(create_date,'%Y-%m-%d %H:%i') 
                order by 1'''

'''
  功能：数据库指标统计语句
'''
st_db_total_num = '''SELECT 
                        DATE_FORMAT(create_date,'%Y-%m-%d %h:%i') AS rq,
                        ROUND(AVG(total_connect),2) AS val
                     FROM `t_monitor_task_db_log`
                       WHERE server_id={}
                         AND  create_date BETWEEN '{} 0:0:0' AND '{} 23:59:59' 
                       GROUP BY DATE_FORMAT(create_date,'%Y-%m-%d %h:%i') 
                     order by 1'''

st_db_active_num = '''SELECT 
                        DATE_FORMAT(create_date,'%Y-%m-%d %h:%i') AS rq,
                        ROUND(AVG(active_connect),2) AS val
                      FROM `t_monitor_task_db_log`
                       WHERE server_id={}
                         AND  create_date BETWEEN '{} 0:0:0' AND '{} 23:59:59' 
                       GROUP BY DATE_FORMAT(create_date,'%Y-%m-%d %h:%i') 
                      order by 1'''


st_db_bak_size = '''SELECT 
                        DATE_FORMAT(create_date,'%Y-%m-%d') AS rq,
                        CASE 
                            WHEN INSTR(total_size,'G')>0 THEN REPLACE(total_size,'G','')*1024
                            WHEN INSTR(total_size,'M')>0 THEN REPLACE(total_size,'M','')
                            ELSE REPLACE(total_size,'K','')/1024
                        END AS val
                    FROM `t_db_backup_total`
                       WHERE db_tag=(SELECT db_tag FROM t_db_config WHERE db_id={}) 
                         AND  create_date BETWEEN '{} 0:0:0' AND '{} 23:59:59' 
                   order by 1'''

st_db_bak_time = '''SELECT 
                        DATE_FORMAT(create_date,'%Y-%m-%d') AS rq,
                        elaspsed_backup AS val
                    FROM `t_db_backup_total`
                       WHERE db_tag=(SELECT db_tag FROM t_db_config WHERE db_id={}) 
                         AND  create_date BETWEEN '{} 0:0:0' AND '{} 23:59:59' 
                    order by 1'''


'''  
   功能：获取mysql连接，以字典返回
'''
def get_ds_mysql_dict(ip,port,service ,user,password):
    conn = pymysql.connect(host=ip, port=int(port),
                           user=user, passwd=password,
                           db=service,charset='utf8',
                           cursorclass = pymysql.cursors.DictCursor)
    return conn


'''
  功能：打印字典
'''
def print_dict(config):
    print('-'.ljust(85,'-'))
    print(' '.ljust(3,' ')+"name".ljust(20,' ')+'value')
    print('-'.ljust(85,'-'))
    for key in config:
      print(' '.ljust(3,' ')+key.ljust(20,' ')+'=',config[key])
    print('-'.ljust(85,'-'))

'''
  功能：根据st语句返回结果集cpu,mem
'''
def get_result(config,st):
    db = config['db_mysql_dict']
    cr = db.cursor()
    cr.execute(st)
    rs = cr.fetchall()
    x=[]
    y=[]
    for i in rs:
        x.append(i['rq'])
        y.append(i['val'])
    return str(x),str(y)

'''
  功能：根据st语句返回结果集db_sz,db_tm
'''
def get_result_pie(config,st):
    db = config['db_mysql_dict']
    cr = db.cursor()
    cr.execute(st)
    rs = cr.fetchall()
    v  = ''
    js = json.dumps(rs)


'''
  功能：根据st语句返回结果集server_info
'''
def get_result_server(config,st):
    db = config['db_mysql_dict']
    cr = db.cursor()
    cr.execute(st)
    rs = cr.fetchall()
    return rs


'''
  功能：根据st语句返回结果集disk_usage
'''
def get_disk_usage_result(config,st):
    db   = config['db_mysql_dict']
    cr   = db.cursor()
    cr.execute(st)
    rs   = cr.fetchall()
    disk = {}
    for d in rs:
        disk[d['rq']]=json.loads(d['val'])
    x = []
    y = []
    for d in disk:
      x.append(d)
      y.append(disk[d])
    return str(x),y

'''
  功能：生成echarts画图options
'''
def gen_options(p_title,p_legend,x_data,y_series):
    v = '''{
                title: {
                    text: 'p_title'
                },
                tooltip: { 
                    trigger: 'axis',
                    axisPointer: {
                        type: 'cross',
                        label: {
                            backgroundColor: '#6a7985'
                        }
                      }
                },
                toolbox: {
                　　show: true,
                　　feature: {
                　　　　saveAsImage: {
                　　　　show:true,
                　　　　excludeComponents :['toolbox'],
                　　　　pixelRatio: 2
                　　　　}
                　　}
                },
                legend: {
                    data:['p_legend']
                },
                xAxis: {
                    data: x_data
                },
                yAxis: {
                    type: 'value'
                },
                series: y_series
        }'''
    v = v.replace('p_title' , p_title)
    v = v.replace('p_legend', p_legend)
    v = v.replace('x_data'  , x_data)
    v = v.replace('y_series', y_series)
    return v


'''
  功能：生成echarts画图series(cpu,mem)
'''
def gen_series(p_title,y_data,p_color):
    t = []
    y = '''{ name: 'p_title', type: 'p_line',areaStyle:{normal:{color:'p_color'}},itemStyle:{normal:{color:'p_color',lineStyle:{color:'p_color'}}},data: y_data}'''
    y = y.replace('p_title',p_title).replace('y_data',y_data).replace('p_line',config['shape']).replace('p_color',p_color)
    t.append(y)
    return str(t).replace('"','')

'''
  功能：生成echarts画图series(db_sz,db_tm)
'''
def gen_series_bar(p_title,y_data,p_color):
    t = []
    y = '''{ name: 'p_title', type: 'bar',areaStyle:{normal:{color:'p_color'}},itemStyle:{normal:{color:'p_color',lineStyle:{color:'p_color'}}},data: y_data}'''
    y = y.replace('p_title',p_title).replace('y_data',y_data).replace('p_line',config['shape']).replace('p_color',p_color)
    t.append(y)
    return str(t).replace('"','')


'''
  功能：生成echarts画图series(disk_usage)
'''
def gen_series_disk_usage(y_data,p_color):
    # 获取图例列表
    h = []
    for data in y_data:
        for ckey in data:
           h.append(ckey)
        break

    # 处理传入字典为一个实单的列表
    t = []
    for data in y_data:
        v = ''
        for ckey in data:
           v = v+ str(data[ckey])+'@'
        t.append(v[0:-1])

    # 动态初始化字典元素为列表
    d = {}
    for v in t:
        for i in range(len(v.split('@'))):
            d[i] = []
        break

    # 写入字典
    for v in t:
        for i in range(len(v.split('@'))):
           d[i].append(float(v.split('@')[i]))

    # 生成series字符串
    s = []
    for i in range(len(h)):
        y = '''{ name: 'p_title', type: 'p_line',areaStyle:{normal:{color:'p_color'}},itemStyle:{normal:{color:'p_color',lineStyle:{color:'p_color'}}}, data: y_data}'''
        y = y.replace('p_title',h[i]).replace('y_data',str(d[i])).replace('p_line',config['shape_disk_uage']).replace('p_color',p_color.split(',')[i])
        s.append(y)
    s = str(s).replace('"','')
    return s,"','".join(h)


'''
 功能：生成html
'''
def gen_html(cfg,html_templete):

    # 更新模板
    html_templete = html_templete.replace('$$HEADER$$', get_server(cfg)['server_desc'])
    html_templete = html_templete.replace('$$TJRQQ$$', config['start_date'])
    html_templete = html_templete.replace('$$TJRQZ$$', config['end_date'])
    html_templete = html_templete.replace('$$BZRQ$$', get_bzrq())

    # cpu使用率
    x_cpu, y_cpu               = get_result(config, st_cpu.format(cfg['server_id'],cfg['start_date'],cfg['end_date']))
    cpu_series                 = gen_series(config['cpu_title'],y_cpu,config['cpu_color'])
    cpu_option                 = gen_options(config['cpu_title'],config['cpu_title'],x_cpu,cpu_series)

    # 更新服务器图表
    html_templete = html_templete.replace('$cpu_option', cpu_option)


    return html_templete

def write_html(cfg):
    filename = '商管-{}-运行周报.html'.format(get_server(cfg)['server_desc'])
    print('write html: {} ok'.format(filename))
    with open(filename, 'w') as f:
        v=gen_html(cfg,html_templete)
        f.write(v)


def set_sytle(doc):
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(10.5)
    doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    return doc

def set_margin(doc):
    section = doc.sections[0]
    print('默认页面的宽度和高度：', section.page_width.cm, section.page_height.cm)  # 打印默认页面宽度和高度
    print("页面上边距：", section.top_margin.cm)
    print("页面下边距：", section.bottom_margin.cm)
    print("页面左边距：", section.left_margin.cm)
    print("页面右边距：", section.right_margin.cm)
    section.top_margin = Cm(1.54)
    section.bottom_margin = Cm(1.54)
    section.left_margin = Cm(1.175)
    section.right_margin = Cm(1.175)
    return doc

def  write_doc(cfg):
     doc = cfg['doc']
     doc = set_sytle(doc)
     doc = set_margin(doc)

     doc.add_heading('商管-系统运行周报V2.0', 0)
     doc.add_heading('一、备份服务器', level=1)
     doc.add_heading('1.1、运行情况概述',level=2)
     p = doc.add_paragraph('\n服务器地址:{}'.format(get_server(cfg)['server_ip']))

     table = doc.add_table(rows=1, cols=6,style='Normal Table')
     hdr_cells         = table.rows[0].cells
     hdr_cells[0].text = '指标'
     hdr_cells[1].text = '最小值'
     hdr_cells[2].text = '最大值'
     hdr_cells[3].text = '平均值'
     hdr_cells[4].text = '告值阀值'
     hdr_cells[5].text = '备注'
     for r in get_result_server(config,
                                st_server.format(config['server_id'], config['start_date'],
                                                 config['end_date'],
                                                 config['server_id'], config['start_date'],
                                                 config['end_date'],
                                                 config['server_id'], config['start_date'],
                                                 config['end_date'])):
         row_cells = table.add_row().cells
         row_cells[0].text = r['flag']
         row_cells[1].text = r['val_min']
         row_cells[2].text = r['val_max']
         row_cells[3].text = r['val_avg']
         row_cells[4].text = r['threshold']
         row_cells[5].text = r['message']

     # doc.add_page_break()
     doc.add_heading('1.2、服务器运行情况', level=2)
     cfg['doc'].save('商管-系统运行周报V2.0.docx')


def stats():
    cfg = get_cfg(config)
    os.system('rm 商管-系统运行周报V2.0.docx')
    doc = Document()
    for s in tj_server:
        print(s['server_id'], s['db_id'])
        cfg['server_id'] = s['server_id']
        cfg['db_id'] = s['db_id']
        cfg['doc']   = doc
        write_doc(cfg)


if __name__ == "__main__":
    stats()


