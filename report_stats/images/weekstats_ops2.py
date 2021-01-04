#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time : 2020/11/23 9:56
# @Author : 马飞
# @File : weekstats_ops.py.py
# @Software: PyCharm

import datetime
import pymysql
import base64
import json
import sys
from pyecharts.charts import Bar
from pyecharts.charts import Line
from pyecharts.render import make_snapshot
from pyecharts import options as opts
from snapshot_phantomjs import snapshot
from pyecharts.globals import ThemeType

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared  import Pt
from docx.shared  import Inches
from docx.oxml.ns import nsdecls
from docx.oxml    import parse_xml


config = {
   'db_string'        : '10.2.39.17:23306:puppet:puppet:Puppet@123',
   'start_date'       : (datetime.date.today() + datetime.timedelta(days = -5)).strftime("%Y-%m-%d"),
   'end_date'         : datetime.date.today().strftime("%Y-%m-%d"),
   'img_width'        : '800px',
   'img_height'       : '250px',
   'tab_title_color'  : '#D3D3D3',
   'theme_type'       : ThemeType.MACARONS
}

server =[
    {  'name':'bi系统',
       'list':[
               {
                    'type': 'db',
                    'list':
                         [
                           {
                               'server_id': 79,
                               'db_id': 137,
                           },
                         ]
                },
                {
                   'type': 'app',
                   'list':
                       [
                           {
                             'server_id': 82,
                             'db_id': '',
                           },
                       ]
               }
        ],
    }
 ]

def get_ds_mysql(ip, port, service, user, password):
    conn = pymysql.connect(host=ip, port=int(port),user=user, passwd=password,db=service, charset='utf8',
                           cursorclass=pymysql.cursors.DictCursor)
    return conn

def print_dict(config):
    print('-'.ljust(85, '-'))
    print(' '.ljust(3, ' ') + "name".ljust(20, ' ') + 'value')
    print('-'.ljust(85, '-'))
    for key in config:
        print(' '.ljust(3, ' ') + key.ljust(20, ' ') + '=', config[key])
    print('-'.ljust(85, '-'))

def get_cfg(cfg):
    db_ip      = cfg['db_string'].split(':')[0]
    db_port    = cfg['db_string'].split(':')[1]
    db_service = cfg['db_string'].split(':')[2]
    db_user    = cfg['db_string'].split(':')[3]
    db_pass    = cfg['db_string'].split(':')[4]
    cfg['db_mysql_dict'] = get_ds_mysql(db_ip, db_port, db_service, db_user, db_pass)
    return cfg

def get_base64():
    with open("line.png", 'rb') as f:
        base64_data = base64.b64encode(f.read())
        s = base64_data.decode()
    return s

def save_image(cfg,item):
    db = cfg['db_mysql_dict']
    cr = db.cursor()
    st = """insert t_db_weekly_data(item_type,item_code,create_date,img_base64) values('{}','{}',now(),'{}')""".\
         format(item['item_type'], item['item_code'], get_base64())
    cr.execute(st)
    cr.close()

def get_server(cfg):
    db = cfg['db_mysql_dict']
    cr = db.cursor()
    cr.execute('select * from t_server where id={}'.format(cfg['server_id']))
    rs = cr.fetchone()
    return rs

def get_db(cfg):
    db = cfg['db_mysql_dict']
    cr = db.cursor()
    cr.execute('select * from t_db_source where id={}'.format(cfg['db_id']))
    rs = cr.fetchone()
    return rs

def get_tj_item(config,p_type):
    db = config['db_mysql_dict']
    cr = db.cursor()
    cr.execute("""select item_type,
                         item_code,
                         item_desc,
                         item_tjsql
                  from t_db_weekly_items 
                  where item_type='{}' order by id
               """.format(p_type))
    rs = cr.fetchall()
    cr.close()
    return rs

def get_xaxis_yaris(p_rs):
    x = []
    y = []
    for r in p_rs:
      x.append(r['rq'])
      y.append(r['val'])
    return x,y

def get_tj_data(cfg,p_type,p_tjsql):
    db = config['db_mysql_dict']
    cr = db.cursor()
    if p_type == 'server':
       cr.execute(p_tjsql.format(cfg['server_id'],cfg['start_date'],cfg['end_date']))
    else:
       cr.execute(p_tjsql.format(cfg['db_id'], cfg['start_date'], cfg['end_date']))
    rs = cr.fetchall()
    cr.close()
    return get_xaxis_yaris(rs)

def get_tj_data_running(cfg,p_type,p_tjsql):
    db = config['db_mysql_dict']
    cr = db.cursor()
    if p_type in('db_detail','server_detail'):
       cr.execute(p_tjsql.replace('$$SERVER_ID$$',str(cfg['server_id']))
                         .replace('$$DB_ID$$', str(cfg['db_id']))
                         .replace('$$TJRQQ$$', cfg['start_date'])
                         .replace('$$TJRQZ$$', cfg['end_date']))
    rs = cr.fetchall()
    cr.close()
    return rs

def get_max_min_val(p_vals):
    print(p_vals)
    keys ={}
    for r in p_vals[0:1]:
        for k in r:
            keys[k]=0
    print('keys=',keys)
    for k in keys:
        n_max = 0
        for r in p_vals:
            t = json.loads(r)
            if float(t[k])>n_max:
                n_max = float(t[k])
        keys[k] = n_max
    return json.dumps(keys)

def get_tj_data_json(cfg,p_tjsql):
    db = config['db_mysql_dict']
    cr = db.cursor()
    cr.execute(p_tjsql.format(cfg['server_id'],cfg['start_date'],cfg['end_date']))
    rs = cr.fetchall()
    print('rs=',rs[0])
    vals  = []
    for r in rs:
      keys = {}
      keys['rq']  = r['rq']
      keys['val'] = []
      vals.append(keys)
    print('vals=',vals)

    for val in vals:
       tmp = []
       for r in rs:
          if val['rq'] == r['rq']:
             tmp.append(r['val'])
       val['val'] = tmp

    for r in vals:
        v = get_max_min_val(r['val'])
        print(r['rq'],v)
    cr.close()
    return get_xaxis_yaris(rs)

def get_cpu_mem_line(cfg,item,x_data,y_data):
    Grapth = (
        Line(init_opts=opts.InitOpts(theme=cfg['theme_type'],width=cfg['img_width'],height=cfg['img_height']))
            .set_global_opts(
            tooltip_opts=opts.TooltipOpts(is_show=False),
            xaxis_opts=opts.AxisOpts(
                type_="category",
                axistick_opts=opts.AxisTickOpts(is_align_with_label=True),
                is_scale=False,
                boundary_gap=False,
             ),
            yaxis_opts=opts.AxisOpts(
                type_="value",
                axistick_opts=opts.AxisTickOpts(is_show=True),
                splitline_opts=opts.SplitLineOpts(is_show=False),
                max_=100,
                min_=0,
             ),
            )
            .set_series_opts(
            areastyle_opts=opts.AreaStyleOpts(opacity=0.5),
            label_opts=opts.LabelOpts(is_show=False),
             )
            .add_xaxis(xaxis_data=x_data)
            .add_yaxis(series_name="{}".format(item['item_desc']),
                       y_axis=y_data,
                       symbol="emptyCircle",
                       is_symbol_show=True,
                       label_opts=opts.LabelOpts(is_show=False),
                       areastyle_opts=opts.AreaStyleOpts(opacity=0.5),
                       is_smooth=True,
                       markpoint_opts=opts.MarkPointOpts(
                           data=[
                               opts.MarkPointItem(type_="max", name="最大值"),
                               opts.MarkPointItem(type_="min", name="最小值"),
                           ]
                       ),
                       markline_opts=opts.MarkLineOpts(
                           data=[opts.MarkLineItem(type_="average", name="平均值")]
                       ),
            )

    )
    return Grapth

def get_disk_net_rw(cfg,item,x_data,y_data):
    Grapth = (
        Line(init_opts=opts.InitOpts(theme=cfg['theme_type'],width=cfg['img_width'],height=cfg['img_height']))
            .set_global_opts(
            tooltip_opts=opts.TooltipOpts(is_show=False),
            xaxis_opts=opts.AxisOpts(
                type_="category",
                axistick_opts=opts.AxisTickOpts(is_align_with_label=True),
                is_scale=False,
                boundary_gap=False,
             ),
            yaxis_opts=opts.AxisOpts(
                type_="value",
                axistick_opts=opts.AxisTickOpts(is_show=True),
                splitline_opts=opts.SplitLineOpts(is_show=False),
             ),
            )
            .set_series_opts(
            areastyle_opts=opts.AreaStyleOpts(opacity=0.5),
            label_opts=opts.LabelOpts(is_show=False),
             )
            .add_xaxis(xaxis_data=x_data)
            .add_yaxis(series_name="{}".format(item['item_desc']),
                       y_axis=y_data,
                       symbol="emptyCircle",
                       is_symbol_show=True,
                       label_opts=opts.LabelOpts(is_show=False),
                       areastyle_opts=opts.AreaStyleOpts(opacity=0.5),
                       is_smooth=True,
                       markpoint_opts=opts.MarkPointOpts(
                           data=[
                               opts.MarkPointItem(type_="max", name="最大值"),
                               opts.MarkPointItem(type_="min", name="最小值"),
                           ]
                       ),
                       markline_opts=opts.MarkLineOpts(
                           data=[opts.MarkLineItem(type_="average", name="平均值")]
                       ),
            )

    )
    return Grapth

def get_disk_usage(cfg,item,x_data,y_data):
    line = Line(init_opts=opts.InitOpts(theme=cfg['theme_type'], width=cfg['img_width'], height=cfg['img_height']))
    line.set_global_opts(
        title_opts=opts.TitleOpts(title=item['item_desc'],
                                  pos_left="center",
                                  title_textstyle_opts=opts.TextStyleOpts(font_size=10)),

        tooltip_opts=opts.TooltipOpts(is_show=False),
        xaxis_opts=opts.AxisOpts(
            type_="category",
            axisline_opts=opts.AxisLineOpts(is_on_zero=True),
            is_scale=False,
            boundary_gap=False,
            axislabel_opts=opts.LabelOpts(rotate=-15,font_size=10),

        ),
        yaxis_opts=opts.AxisOpts(
            type_="value",
            axistick_opts=opts.AxisTickOpts(is_show=True),
            splitline_opts=opts.SplitLineOpts(is_show=False),
            max_=100,
            min_=0
        ),
        legend_opts=opts.LegendOpts(pos_left="7%",textstyle_opts=opts.TextStyleOpts(font_size=10)),
    )
    line.set_series_opts(
        areastyle_opts=opts.AreaStyleOpts(opacity=0.5),
        label_opts=opts.LabelOpts(is_show=False,font_size=10),
    )

    line.add_xaxis(xaxis_data=x_data)

    #init key
    duk = {}
    for data in y_data[0:2]:
        du = json.loads(data)
        for key in du:
            duk[key] = []

    #load data into duk
    for data in y_data:
        du = json.loads(data)
        for key in du:
            duk[key].append(du[key])

    for key in duk:
        line.add_yaxis(series_name="{}".format(key),
                       y_axis=duk[key],
                       symbol="emptyCircle",
                       symbol_size=8,
                       is_hover_animation=False,
                       is_symbol_show=True,
                       label_opts=opts.LabelOpts(is_show=False,font_size=10),
                       linestyle_opts=opts.LineStyleOpts(width=1.5),
                       areastyle_opts=opts.AreaStyleOpts(opacity=0.5),
                       is_smooth=True,
                       markpoint_opts=opts.MarkPointOpts(
                           data=[
                               opts.MarkPointItem(type_="max", name="最大值"),
                               opts.MarkPointItem(type_="min", name="最小值"),
                           ]
                       ),
                       markline_opts=opts.MarkLineOpts(
                           data=[opts.MarkLineItem(type_="average", name="平均值")]
                       ),
         )

    return line

def get_disk_usage_bar(cfg,item,x_data,y_data):
    line = Bar(init_opts=opts.InitOpts(theme=cfg['theme_type'],width=cfg['img_width'],height=cfg['img_height']))
    line.set_global_opts(
            tooltip_opts=opts.TooltipOpts(is_show=False),
            xaxis_opts=opts.AxisOpts(
                axislabel_opts=opts.LabelOpts(rotate=-15)
             ),
            yaxis_opts=opts.AxisOpts(
                type_="value",
             ),
            )
    line.set_series_opts(
            areastyle_opts=opts.AreaStyleOpts(opacity=0.5),
            label_opts=opts.LabelOpts(is_show=False),
           )

    line.add_xaxis(xaxis_data=x_data)
    line.add_yaxis(series_name=item['item_desc'],yaxis_data=y_data)

    return line

def tabBgColor(table,cols,colorStr):
    shading_list = locals()
    for i in range(cols):
        shading_list['shading_elm_'+str(i)] = parse_xml(r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'),bgColor = colorStr))
        table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_list['shading_elm_'+str(i)])

def write_doc_header(cfg):
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run = document.add_heading('', 0).add_run(u'商管-系统运行周报')
    run.font.name = u'黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

    for i in range(18):
        p1 = document.add_paragraph()
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p2 = document.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p2.add_run('编制日期：{}'.format(cfg['start_date']))

    p3 = document.add_paragraph()
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p3.add_run('编制人员：好生活平台组')
    document.add_page_break()
    return document

def write_db_server(svr,document):
    run = document.add_heading('', 1).add_run(u'一、BI系统')
    run.font.name = u'黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

    run = document.add_heading('', 2).add_run(u'1.1、数据库服务器')
    run.font.name = u'黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

    run = document.add_heading('', 3).add_run(u'1.1.1、运行情况概述')
    run.font.name = u'黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

    for svr in server:


        cfg['db_id'] = svr['db_id']
        cfg['server_id'] = svr['server_id']
        cfg['server'] = get_server(cfg)
        cfg['db'] = get_db(cfg)
        for item in get_tj_item(cfg, 'db_detail'):
            res = get_tj_data_running(cfg, 'db_detail', item['item_tjsql'])
            tab_p1 = document.add_paragraph()
            tab_p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            if item['item_code'] == 'server_info':
                tab_p1.add_run(item['item_desc'].format(cfg['server']['server_ip'])).bold = True
            else:
                tab_p1.add_run(item['item_desc'].format(cfg['db']['ip'], cfg['db']['port'])).bold = True

            table = document.add_table(rows=len(res), cols=len(res[0]), style='Table Grid')
            table_run1 = table.cell(0, 0).paragraphs[0].add_run(item['item_desc'])
            table_run1.font.name = '隶书'
            document.styles['Normal'].font.size = Pt(8)
            table_run1.element.rPr.rFonts.set(qn('w:eastAsia'), '隶书')

            i = 0
            for r in res:
                j = 0
                for c in r:
                    if i == 0:
                        table.cell(i, j).text = r[c]
                        tabBgColor(table, len(res[0]), cfg['tab_title_color'])
                    else:
                        table.cell(i, j).text = r[c]
                    j = j + 1
                i = i + 1
            document.add_paragraph()

    pars = document.paragraphs
    for par in pars:
        if par.style.name in ("Heading 1", "Heading 2", "Heading 3"):
            par.paragraph_format.space_before = Pt(10)
            par.paragraph_format.space_after = Pt(10)

    document.add_heading('1.1.2、服务器运行情况', 3)
    for svr in server:
        cfg['server_id'] = svr['server_id']
        for item in get_tj_item(cfg, 'server'):
            if item['item_code'] in ('cpu', 'memory'):
                x_data, y_data = get_tj_data(cfg, 'server', item['item_tjsql'])
                Grapth = get_cpu_mem_line(cfg, item, x_data, y_data)
            elif item['item_code'] in ('disk_read', 'disk_write', 'net_in', 'net_out'):
                x_data, y_data = get_tj_data(cfg, 'server', item['item_tjsql'])
                Grapth = get_disk_net_rw(cfg, item, x_data, y_data)
            elif item['item_code'] in ('disk_usage',):
                x_data, y_data = get_tj_data(cfg, 'server', item['item_tjsql'])
                Grapth = get_disk_usage(cfg, item, x_data, y_data)
            else:
                pass
            make_snapshot(snapshot, Grapth.render(), item['item_code'] + '.png', pixel_ratio=1)
            print(item['item_code'] + '.png generate ok!')
            document.add_picture(item['item_code'] + '.png', width=Inches(6))

    document.add_heading('1.1.3、数据库运行情况', 3)
    for svr in server:
        cfg['db_id'] = svr['db_id']
        for item in get_tj_item(cfg, 'db'):
            if item['item_code'] in ('backup_size', 'backup_time'):
                x_data, y_data = get_tj_data(cfg, 'db', item['item_tjsql'])
                Grapth = get_disk_usage_bar(cfg, item, x_data, y_data)
            else:
                x_data, y_data = get_tj_data(cfg, 'db', item['item_tjsql'])
                if x_data == [] or y_data == []:
                    continue
                Grapth = get_disk_net_rw(cfg, item, x_data, y_data)
            make_snapshot(snapshot, Grapth.render(), item['item_code'] + '.png', pixel_ratio=1)
            print(item['item_code'] + '.png generate ok!')
            document.add_picture(item['item_code'] + '.png', width=Inches(6))

    document.add_heading('1.2、应用服务器', 2)
    document.add_heading('1.2.1、运行情况概述', 3)
    document.add_heading('1.2.2、服务器运行情况', 3)

def stats():
    cfg =  get_cfg(config)
    doc = write_doc_header(cfg)
    for svr in server:
        print(json.dumps(svr, ensure_ascii=False, indent=4, separators=(',', ':')))
        doc = write_db_server(svr,doc)
    doc.save('商管运行周报.docx')

if __name__ == "__main__":
    stats()
